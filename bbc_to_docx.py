import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO

url = "https://www.bbc.com/travel/article/20230313-the-slowest-train-journey-in-india"
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36"
}

response = requests.get(url, headers=headers)
soup = BeautifulSoup(response.text, "html.parser")

doc = Document()

article = soup.find("article")
if not article:
    print("Cannot find article!")
    exit()

title = article.find("h1")
if title:
    doc.add_heading(title.get_text(strip=True), level=0)

# Description
desc = article.find("p", class_="ssrcss-1q0x1qg-Paragraph")
if desc:
    doc.add_paragraph(desc.get_text(strip=True))

# Get content and images
for elem in article.find_all(["h2", "h3", "p", "figure"], recursive=True):
    if elem.get("data-component") in ["tag-list-block", "advertisement-block"]:
        break

    if elem.name in ["h2", "h3"]:
        text = elem.get_text(strip=True)
        if text:
            doc.add_heading(text, level=2)

    elif elem.name == "p":
        text = elem.get_text(strip=True)
        if text:
            doc.add_paragraph(text)

    elif elem.name == "figure":
        img_tag = elem.find("img", {"srcset": True})
        if img_tag:
            srcset = img_tag["srcset"]
            img_url = srcset.split(",")[-1].strip().split(" ")[0]
            print("Ảnh gốc:", img_url)

            # Download image
            img_res = requests.get(img_url)
            image = Image.open(BytesIO(img_res.content))

            # Convert to JPEG for Word
            img_io = BytesIO()
            image.convert("RGB").save(img_io, format="JPEG")
            img_io.seek(0)

            # Insert image into Word
            doc.add_picture(img_io, width=Inches(4))

        caption = elem.find("figcaption")
        if caption:
            para = doc.add_paragraph(caption.get_text(strip=True))
            if para.runs:
                para.runs[0].italic = True

# Save Word file
doc.save("bbc_article.docx")
print("Finished Saving")
